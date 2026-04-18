"""
Blog Writer — Riot PR Desk
Write SEO-optimised blog posts for rioteliquid.com/blogs/news.
"""

import datetime
import io
import json
import re
import streamlit as st

from services.ai_engine import is_configured as ai_configured, generate_stream, refine_text_sync, generate
from services.blog_library import save_blog, get_all_blogs, search_blogs, update_blog_status, delete_blog, add_version
from utils.prompts import BLOG_PROMPT
from utils.styles import apply_global_styles, render_sidebar

# ---------------------------------------------------------------------------
# Page config
# ---------------------------------------------------------------------------

st.set_page_config(page_title="Blog Writer | Riot PR Desk", page_icon="📝", layout="wide")

apply_global_styles()
render_sidebar()

st.title("📝 Blog Writer")
st.markdown(
    "Write SEO-optimised blog posts for rioteliquid.com/blogs/news — "
    "Riot's owned platform for setting the agenda, building search traffic "
    "and speaking directly to readers."
)

st.divider()

# ---------------------------------------------------------------------------
# AI guard
# ---------------------------------------------------------------------------

if not ai_configured():
    st.error(
        "**AI engine not configured.** Add your `ANTHROPIC_API_KEY` or `OPENAI_API_KEY` "
        "to `.env` to get started. Then set `AI_PROVIDER` to `anthropic` or `openai`.",
        icon="🔑",
    )
    st.stop()

# ---------------------------------------------------------------------------
# Status badges (used in both tabs)
# ---------------------------------------------------------------------------

STATUS_BADGES = {"draft": "⚪ Draft", "ready": "✅ Ready", "published": "🌐 Published"}

BLOG_TYPE_OPTIONS = [
    "Industry Commentary",
    "Harm Reduction",
    "Product Education",
    "Campaign Story",
    "Behind The Scenes",
    "News-Jack",
    "Opinion / Thought Leadership",
]
BLOG_TYPE_DESCRIPTIONS = {
    "Industry Commentary": "Riot's take on regulatory news, policy, industry trends",
    "Harm Reduction": "Educational content about switching from smoking, nicotine science",
    "Product Education": "How devices work, flavour guides, choosing the right product",
    "Campaign Story": "Telling the story behind a Riot Activist campaign or stunt",
    "Behind The Scenes": "British manufacturing, company culture, team stories",
    "News-Jack": "Riot's opinion piece pegged to a trending news story",
    "Opinion / Thought Leadership": "Riot's perspective on a big issue",
}


# ---------------------------------------------------------------------------
# Section parser
# ---------------------------------------------------------------------------

def _parse_blog(raw: str) -> dict:
    sections = {}
    current_section = None
    current_content = []
    section_markers = {
        "1. SEO PACKAGE": "SEO Package",
        "2. BLOG POST": "Blog Post",
        "3. IMAGE SUGGESTIONS": "Image Suggestions",
        "4. INTERNAL LINKS": "Internal Links",
        "5. SOCIAL PROMOTION": "Social Promotion",
    }
    for line in raw.split("\n"):
        matched = False
        for marker, label in section_markers.items():
            if marker.lower() in line.lower().replace("#", "").strip().replace("*", ""):
                if current_section:
                    sections[current_section] = "\n".join(current_content).strip()
                current_section = label
                current_content = []
                matched = True
                break
        if not matched and current_section:
            current_content.append(line)
    if current_section:
        sections[current_section] = "\n".join(current_content).strip()
    if not sections:
        sections["Full Response"] = raw
    return sections


# ---------------------------------------------------------------------------
# AI suggestion helper
# ---------------------------------------------------------------------------

def _generate_blog_suggestions(topic: str) -> dict:
    """Ask the AI to recommend blog type, primary keyword and secondary keywords for a topic.

    Returns a dict with keys: blog_type, primary_keyword, secondary_keywords, rationale.
    Returns an empty dict on any failure (UI falls back to manual mode gracefully).
    """
    prompt = f"""You are an SEO and content strategy expert for Riot Labs — a British vape brand based in Norwich.

Riot makes premium e-liquids and vape devices at its UK factory. The brand is activist, direct and irreverent.
Its blog lives at rioteliquid.com/blogs/news and targets: adult vapers, people switching from smoking, and
industry observers. SEO is critical — posts rank on Google for real search queries.

A story has come through the news desk. Analyse it and recommend the optimal blog settings.

STORY / TOPIC:
{topic[:2000]}

Respond with ONLY valid JSON — no markdown fences, no commentary, nothing else:
{{
  "blog_type": "one of: Industry Commentary | Harm Reduction | Product Education | Campaign Story | Behind The Scenes | News-Jack | Opinion / Thought Leadership",
  "primary_keyword": "the single highest-value SEO phrase for this topic (what UK people actually type into Google)",
  "secondary_keywords": ["phrase 2", "phrase 3", "phrase 4", "phrase 5", "phrase 6"],
  "rationale": "1-2 sentences: why this blog type fits, and why these keywords have the best search intent"
}}

Rules:
- Breaking news or trend stories → News-Jack
- Regulatory/policy topics → Industry Commentary
- Educational content about switching → Harm Reduction
- Product features/guides → Product Education
- Primary keyword must be conversational and specific (e.g. "vaping UK law 2025" not "vaping regulatory landscape")
- All keywords should be what real people search on Google UK — long-tail is fine
- Secondary keywords should cover different search intents (e.g. questions, comparisons, local variants)
- Keep rationale under 40 words — punchy and practical"""

    try:
        raw = generate(prompt)
        # Strip any accidental markdown fencing
        clean = re.sub(r"```(?:json)?", "", raw).strip()
        match = re.search(r"\{[\s\S]*\}", clean)
        if match:
            return json.loads(match.group())
        return {}
    except Exception:
        return {}


# ---------------------------------------------------------------------------
# Word export helper
# ---------------------------------------------------------------------------

def _build_blog_docx(sections: dict, blog_title: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title_para = doc.add_heading("RIOT PR DESK — BLOG WRITER (DRAFT)", 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_heading(blog_title or "Blog Post", 1)

    meta = doc.add_paragraph(
        f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}  |  "
        "Status: DRAFT — Requires approval before publishing"
    )
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_paragraph()

    section_order = ["SEO Package", "Blog Post", "Image Suggestions", "Internal Links", "Social Promotion"]
    ordered_keys = [k for k in section_order if k in sections]
    ordered_keys += [k for k in sections if k not in section_order]

    for section_name in ordered_keys:
        content = sections.get(section_name, "")
        if not content.strip():
            continue

        doc.add_page_break()
        doc.add_heading(section_name, 2)

        if section_name == "Blog Post":
            draft_para = doc.add_paragraph("DRAFT — Requires approval before publishing")
            draft_para.runs[0].font.size = Pt(9)
            draft_para.runs[0].font.color.rgb = RGBColor(0xFF, 0x66, 0x00)
            draft_para.runs[0].bold = True
            doc.add_paragraph()

        for block in content.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            if block.startswith("## "):
                doc.add_heading(block[3:].strip(), 3)
            elif block.startswith("# "):
                doc.add_heading(block[2:].strip(), 2)
            elif block.startswith("### "):
                doc.add_heading(block[4:].strip(), 4)
            else:
                doc.add_paragraph(block)

        doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# ---------------------------------------------------------------------------
# Tabs
# ---------------------------------------------------------------------------

tab_write, tab_library = st.tabs(["✍️ Write", "📚 Blog Library"])

# ===========================================================================
# TAB 1 — WRITE
# ===========================================================================

with tab_write:

    # ── Capture load-trigger before any widgets render ───────────────────────
    _trigger_suggest = st.session_state.pop("blog_suggest_on_load", False)

    # ── Step 1: Topic ────────────────────────────────────────────────────────
    st.markdown("### Step 1 — Topic")

    topic = st.text_area(
        "What's the blog about?",
        key="blog_topic",
        height=120,
        placeholder=(
            "e.g. The UK government's disposable vape ban — what it means for adult vapers "
            "who've made the switch..."
        ),
    )

    st.caption("Quick-start:")
    qs_col1, qs_col2, qs_col3, qs_col4 = st.columns(4)
    quickstarts = {
        "Vape Tax Explainer": (
            "The UK government's vaping products duty — what it means for vapers, why the industry "
            "is concerned, and Riot's position on fair taxation of harm reduction products."
        ),
        "Harm Reduction Science": (
            "The science behind vaping as a harm reduction tool — what the evidence says about "
            "switching from smoking, nicotine, and long-term health outcomes."
        ),
        "Switching from Smoking": (
            "A practical guide to switching from cigarettes to vaping — how to choose your first "
            "device, what nicotine strength to start with, and what to expect in the first weeks."
        ),
        "British Manufacturing": (
            "Why Riot Labs makes its products in Britain — quality control, traceability, "
            "supporting UK jobs, and why it matters in an industry flooded with unregulated imports."
        ),
    }
    for col, (label, value) in zip([qs_col1, qs_col2, qs_col3, qs_col4], quickstarts.items()):
        with col:
            if st.button(label, key=f"blog_qs_{label}", use_container_width=True):
                st.session_state["blog_topic"] = value
                for _k in ["blog_suggestions", "blog_suggestions_applied", "blog_suggestions_for_topic"]:
                    st.session_state.pop(_k, None)
                st.rerun()

    # Clear stale suggestions if the topic has changed since they were generated
    _sugg_for = st.session_state.get("blog_suggestions_for_topic", "")
    if "blog_suggestions" in st.session_state and _sugg_for and topic.strip() != _sugg_for:
        for _k in ["blog_suggestions", "blog_suggestions_applied", "blog_suggestions_for_topic"]:
            st.session_state.pop(_k, None)

    # Auto-suggest (triggered by News Desk or manual button below)
    if _trigger_suggest and topic.strip() and "blog_suggestions" not in st.session_state:
        with st.spinner("🤖 Analysing story and suggesting the best blog settings..."):
            _sugg_result = _generate_blog_suggestions(topic.strip())
        st.session_state["blog_suggestions"] = _sugg_result
        st.session_state["blog_suggestions_for_topic"] = topic.strip()

    st.divider()

    # ── Steps 2 & 3: Blog Settings ───────────────────────────────────────────
    _has_suggestions = bool(st.session_state.get("blog_suggestions"))

    if _has_suggestions:
        sugg = st.session_state["blog_suggestions"]

        # Pre-populate widget session state keys from suggestions (once only)
        if not st.session_state.get("blog_suggestions_applied"):
            sugg_type = sugg.get("blog_type", "")
            if sugg_type in BLOG_TYPE_OPTIONS:
                st.session_state["blog_blog_type"] = sugg_type
            pk = sugg.get("primary_keyword", "")
            if pk:
                st.session_state["blog_primary_keyword"] = pk
            sk = sugg.get("secondary_keywords", [])
            if sk:
                st.session_state["blog_secondary_keywords"] = (
                    ", ".join(sk) if isinstance(sk, list) else str(sk)
                )
            st.session_state["blog_suggestions_applied"] = True

        # AI rationale callout
        if sugg.get("rationale"):
            st.info(f"🤖 **AI Analysis:** {sugg['rationale']}")

        st.markdown("### 🎯 Blog Settings — AI Suggested")
        st.caption(
            "The AI has analysed the story and pre-filled these settings. "
            "Review each field and edit anything before generating."
        )

        smart_c1, smart_c2 = st.columns(2)
        with smart_c1:
            blog_type = st.selectbox(
                "Blog type",
                options=BLOG_TYPE_OPTIONS,
                format_func=lambda t: f"{t} — {BLOG_TYPE_DESCRIPTIONS[t]}",
                key="blog_blog_type",
            )
        with smart_c2:
            primary_keyword = st.text_input(
                "Primary keyword",
                key="blog_primary_keyword",
                placeholder="e.g. disposable vape ban UK",
            )

        secondary_keywords_input = st.text_input(
            "Secondary keywords (comma-separated)",
            key="blog_secondary_keywords",
            placeholder="e.g. vaping regulations, e-cigarette, quit smoking",
        )
        st.caption(
            "💡 Use terms people actually search — "
            "'how to quit smoking with a vape', not 'cessation journey'."
        )

        reanalyse_c, manual_c = st.columns(2)
        with reanalyse_c:
            if topic.strip() and st.button(
                "🔄 Re-analyse story", use_container_width=True, key="blog_reanalyse"
            ):
                for _k in ["blog_suggestions", "blog_suggestions_applied", "blog_suggestions_for_topic"]:
                    st.session_state.pop(_k, None)
                st.session_state["blog_suggest_on_load"] = True
                st.rerun()
        with manual_c:
            if st.button("✏️ Switch to manual mode", use_container_width=True, key="blog_to_manual"):
                for _k in ["blog_suggestions", "blog_suggestions_applied", "blog_suggestions_for_topic"]:
                    st.session_state.pop(_k, None)
                st.rerun()

    else:
        # Manual mode — no AI suggestions yet
        if topic.strip():
            sugg_btn_col, _ = st.columns([2, 3])
            with sugg_btn_col:
                if st.button(
                    "🤖 Suggest blog type & keywords →",
                    type="secondary",
                    use_container_width=True,
                    key="blog_manual_suggest_btn",
                ):
                    st.session_state["blog_suggest_on_load"] = True
                    st.rerun()

        st.markdown("### Step 2 — Blog Type")
        blog_type = st.selectbox(
            "Blog type",
            options=BLOG_TYPE_OPTIONS,
            format_func=lambda t: f"{t} — {BLOG_TYPE_DESCRIPTIONS[t]}",
            key="blog_blog_type",
        )

        st.divider()

        st.markdown("### Step 3 — SEO Keywords")
        primary_keyword = st.text_input(
            "Primary keyword",
            key="blog_primary_keyword",
            placeholder="e.g. disposable vape ban UK",
        )
        secondary_keywords_input = st.text_input(
            "Secondary keywords (comma-separated)",
            key="blog_secondary_keywords",
            placeholder="e.g. vaping regulations, e-cigarette, quit smoking",
        )
        st.caption(
            "💡 Use terms people actually search — "
            "'how to stop smoking with a vape', not 'cessation journey'."
        )

    st.divider()

    # ── Tone & Length ─────────────────────────────────────────────────────────
    _tone_step_label = "### Step 2 — Tone & Length" if _has_suggestions else "### Step 4 — Tone & Length"
    st.markdown(_tone_step_label)

    tone_col, length_col = st.columns(2)
    with tone_col:
        tone_dial = st.select_slider(
            "Tone",
            options=["Very measured", "Measured", "Balanced", "Bold", "Very bold"],
            value="Balanced",
            key="blog_tone_dial",
        )
    with length_col:
        word_count_option = st.selectbox(
            "Length",
            options=["Short (~600 words)", "Standard (~1000 words)", "Long (~1500 words)"],
            index=1,
            key="blog_word_count_option",
        )

    word_count_map = {
        "Short (~600 words)": "approximately 600 words",
        "Standard (~1000 words)": "approximately 1000 words",
        "Long (~1500 words)": "approximately 1500 words",
    }
    word_count = word_count_map[word_count_option]

    st.divider()

    # ── Generate button ──────────────────────────────────────────────────────
    generate_clicked = st.button(
        "📝 Write Blog Post",
        type="primary",
        use_container_width=True,
        disabled=not (topic.strip() and primary_keyword.strip()),
        key="blog_generate_btn",
    )

    # ── Generation logic ─────────────────────────────────────────────────────
    if generate_clicked and topic.strip() and primary_keyword.strip():
        secondary_kw_list = [kw.strip() for kw in secondary_keywords_input.split(",") if kw.strip()]
        secondary_kw_str = ", ".join(secondary_kw_list) if secondary_kw_list else "none specified"

        prompt = BLOG_PROMPT.format(
            topic=topic.strip(),
            blog_type=blog_type,
            primary_keyword=primary_keyword.strip(),
            secondary_keywords=secondary_kw_str,
            tone_dial=tone_dial,
            word_count=word_count,
        )

        try:
            with st.status("✍️ Writing your blog post...", expanded=True) as status:
                st.write(
                    f"Crafting a {word_count_option.lower()} {blog_type.lower()} post "
                    f"optimised for '{primary_keyword}'..."
                )
                stream = generate_stream(prompt)
                raw_response = st.write_stream(stream)
                status.update(label="✅ Blog post ready!", state="complete", expanded=False)

            # Store raw + inputs
            st.session_state["blog_raw"] = raw_response
            st.session_state["blog_stored_topic"] = topic.strip()
            st.session_state["blog_stored_type"] = blog_type
            st.session_state["blog_stored_primary_keyword"] = primary_keyword.strip()
            st.session_state["blog_stored_secondary_keywords"] = secondary_kw_list
            st.session_state["blog_stored_tone"] = tone_dial
            st.session_state["blog_stored_word_count"] = word_count_option

            # Parse into sections
            sections = _parse_blog(raw_response)
            st.session_state["blog_sections"] = sections

            # Auto-save to blog library (silent fail)
            try:
                saved = save_blog(
                    topic=topic.strip(),
                    sections=sections,
                    blog_type=blog_type,
                    primary_keyword=primary_keyword.strip(),
                    secondary_keywords=secondary_kw_list,
                )
                st.session_state["blog_saved_id"] = saved["id"]
            except Exception:
                pass

            st.rerun()

        except Exception as e:
            st.error(f"Generation failed: {e}")

    # ── Display sections ──────────────────────────────────────────────────────
    if "blog_sections" in st.session_state:
        sections = st.session_state["blog_sections"]
        stored_topic = st.session_state.get("blog_stored_topic", "")
        stored_primary_kw = st.session_state.get("blog_stored_primary_keyword", primary_keyword.strip())

        st.divider()
        st.markdown("### 📄 Your Blog Post")
        st.caption("DRAFT — All outputs require approval before publishing.")

        for section_name, content in sections.items():
            with st.expander(f"**{section_name}**", expanded=True):
                st.markdown(content)

                if section_name == "Blog Post":
                    st.caption(f"📊 Word count: ~{len(content.split())} words")

                st.code(content, language=None)

                safe_key = section_name.replace(" ", "_").replace("/", "_")
                st.divider()
                action_col1, action_col2 = st.columns(2)

                # ── Regenerate ────────────────────────────────────────────────
                with action_col1:
                    regen_key = f"blog_regen_{safe_key}"
                    if st.button(
                        "🔄 Regenerate section",
                        key=regen_key,
                        help="Rewrite this section from scratch with a fresh approach",
                        use_container_width=True,
                    ):
                        with st.spinner(f"Regenerating {section_name}..."):
                            try:
                                regen_prompt = (
                                    f"You previously generated a blog post for Riot Labs. "
                                    f"Please regenerate ONLY the '{section_name}' section "
                                    f"with a fresh approach. Keep the same facts and key messages "
                                    f"but use different wording, structure or angle. "
                                    f"Do not include any other sections.\n\n"
                                    f"Original topic: {stored_topic[:400]}\n\n"
                                    f"Current {section_name}:\n{content}\n\n"
                                    f"Generate an improved version of just the {section_name}:"
                                )
                                new_content = generate(regen_prompt)

                                # Save version before replacing
                                saved_id = st.session_state.get("blog_saved_id")
                                if saved_id:
                                    try:
                                        add_version(
                                            saved_id,
                                            dict(st.session_state["blog_sections"]),
                                            note=f"Before regenerating {section_name}",
                                        )
                                    except Exception:
                                        pass

                                updated = dict(st.session_state["blog_sections"])
                                updated[section_name] = new_content
                                st.session_state["blog_sections"] = updated
                                st.rerun()
                            except Exception as e:
                                st.error(f"Regeneration failed: {e}")

                # ── Refine with AI ────────────────────────────────────────────
                with action_col2:
                    refine_toggle_key = f"blog_refine_active_{safe_key}"
                    if st.button(
                        "✏️ Refine with AI" if not st.session_state.get(refine_toggle_key) else "✕ Close editor",
                        key=f"blog_refine_toggle_{safe_key}",
                        help="Give the AI a specific instruction to improve this section",
                        use_container_width=True,
                    ):
                        st.session_state[refine_toggle_key] = not st.session_state.get(refine_toggle_key, False)
                        st.rerun()

                # ── Inline AI editing panel ───────────────────────────────────
                if st.session_state.get(f"blog_refine_active_{safe_key}"):
                    st.markdown(
                        """<div style="background:#111;border:1px solid #E8192C33;border-radius:4px;"""
                        """padding:0.75rem 1rem 0.5rem 1rem;margin-top:0.5rem">""",
                        unsafe_allow_html=True,
                    )

                    quick_prompts = [
                        "Make it shorter and punchier",
                        "Make the tone more provocative",
                        "Make the tone more measured",
                        "Add a stronger opening line",
                        "Improve SEO keyword density",
                        "Remove any corporate clichés",
                        "Add a specific statistic",
                        "Strengthen the call to action",
                    ]
                    st.caption("✏️ **AI EDIT** — Tell the AI exactly what to change:")

                    qp_col1, qp_col2 = st.columns([3, 1])
                    with qp_col1:
                        refine_instruction = st.text_input(
                            "Instruction",
                            key=f"blog_refine_instruction_{safe_key}",
                            placeholder="e.g. 'Make it punchier', 'Add a stat', 'Shorten by half'...",
                            label_visibility="collapsed",
                        )
                    with qp_col2:
                        apply_refine = st.button(
                            "Apply →",
                            key=f"blog_refine_apply_{safe_key}",
                            type="primary",
                            use_container_width=True,
                            disabled=not refine_instruction.strip(),
                        )

                    # Quick-prompt chips
                    chip_cols = st.columns(4)
                    for ci, qp in enumerate(quick_prompts):
                        with chip_cols[ci % 4]:
                            if st.button(qp, key=f"blog_qp_{safe_key}_{ci}", use_container_width=True):
                                st.session_state[f"blog_refine_instruction_{safe_key}"] = qp
                                st.session_state[f"blog_qp_trigger_{safe_key}"] = True
                                st.rerun()

                    st.markdown("</div>", unsafe_allow_html=True)

                    should_apply = apply_refine or st.session_state.pop(f"blog_qp_trigger_{safe_key}", False)
                    instruction_to_use = st.session_state.get(f"blog_refine_instruction_{safe_key}", "").strip()

                    if should_apply and instruction_to_use:
                        # Save version before editing
                        saved_id = st.session_state.get("blog_saved_id")
                        if saved_id:
                            try:
                                add_version(
                                    saved_id,
                                    dict(st.session_state["blog_sections"]),
                                    note=f"Before AI edit: '{instruction_to_use}' on {section_name}",
                                )
                            except Exception:
                                pass

                        with st.spinner(f"Applying: \"{instruction_to_use}\"..."):
                            try:
                                new_content = refine_text_sync(content, instruction_to_use, context=section_name)
                                updated = dict(st.session_state["blog_sections"])
                                updated[section_name] = new_content
                                st.session_state["blog_sections"] = updated
                                st.session_state[f"blog_refine_active_{safe_key}"] = False
                                st.rerun()
                            except Exception as e:
                                st.error(f"Refinement failed: {e}")

        # ── Publishing readiness checklist ────────────────────────────────────
        st.divider()

        checklist_items = []

        kw_in_seo = stored_primary_kw.lower() in sections.get("SEO Package", "").lower()
        checklist_items.append(("Primary keyword in SEO Package title", kw_in_seo))

        has_meta = "meta description" in sections.get("SEO Package", "").lower()
        checklist_items.append(("Meta description present", has_meta))

        h2_count = sections.get("Blog Post", "").count("## ")
        has_h2s = h2_count >= 3
        checklist_items.append(("3+ H2 headings in blog post", has_h2s))

        has_images = len(sections.get("Image Suggestions", "")) > 50
        checklist_items.append(("Image suggestions present", has_images))

        has_cta = any(
            w in sections.get("Blog Post", "").lower()
            for w in ["click", "visit", "shop", "discover", "read more", "explore", "find out"]
        )
        checklist_items.append(("CTA present in blog post", has_cta))

        blog_words = len(sections.get("Blog Post", "").split())
        has_word_count = blog_words >= 800
        checklist_items.append(("Word count 800+", has_word_count))

        passed = sum(1 for _, v in checklist_items if v)
        total = len(checklist_items)
        score_pct = round(passed / total * 100)
        score_color = "🟢" if score_pct >= 80 else "🟡" if score_pct >= 60 else "🔴"

        with st.expander(
            f"{score_color} Publishing Readiness: {score_pct}% ({passed}/{total} checks passed)",
            expanded=score_pct < 80,
        ):
            for label, passed_check in checklist_items:
                icon = "✅" if passed_check else "⚠️"
                st.caption(f"{icon} {label}")
            if score_pct == 100:
                st.success("All checks passed. This post is ready for review.")

        # ── Export ────────────────────────────────────────────────────────────
        st.divider()
        st.markdown("### Export")

        export_col1, export_col2, export_col3 = st.columns(3)

        # Build full plain text export
        full_text = (
            f"RIOT PR DESK — BLOG WRITER (DRAFT)\n"
            f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
            f"{'=' * 60}\n\n"
        )
        full_text += "\n\n".join(
            f"{'=' * 60}\n{name}\n{'=' * 60}\n\n{content}"
            for name, content in sections.items()
        )

        with export_col1:
            st.download_button(
                "📥 Download .txt",
                data=full_text,
                file_name=f"riot_blog_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                mime="text/plain",
                use_container_width=True,
            )

        with export_col2:
            try:
                # Derive blog title from SEO Package if possible
                seo_text = sections.get("SEO Package", "")
                blog_title = ""
                for line in seo_text.splitlines():
                    ls = line.strip()
                    if ls.lower().startswith("**title tag:**"):
                        blog_title = ls[len("**title tag:**"):].strip().strip("[]")
                        break
                    if ls.lower().startswith("title tag:"):
                        blog_title = ls[len("title tag:"):].strip().strip("[]")
                        break
                if not blog_title:
                    blog_title = stored_topic[:80]

                docx_bytes = _build_blog_docx(sections, blog_title)
                st.download_button(
                    "📄 Download .docx",
                    data=docx_bytes,
                    file_name=f"riot_blog_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except ImportError:
                st.caption("Install python-docx for Word export.")

        with export_col3:
            if "blog_saved_id" in st.session_state:
                st.caption("✅ Auto-saved to Blog Library")


# ===========================================================================
# TAB 2 — BLOG LIBRARY
# ===========================================================================

with tab_library:
    st.markdown("### 📚 Blog Library")

    # ── Stats row ─────────────────────────────────────────────────────────────
    all_blogs = get_all_blogs()
    total_blogs = len(all_blogs)
    published_count = sum(1 for b in all_blogs if b.get("status") == "published")
    draft_count = sum(1 for b in all_blogs if b.get("status") == "draft")
    avg_words = (
        round(sum(b.get("word_count", 0) for b in all_blogs) / total_blogs)
        if total_blogs > 0
        else 0
    )

    stat_c1, stat_c2, stat_c3, stat_c4 = st.columns(4)
    with stat_c1:
        st.metric("Total Blogs", total_blogs)
    with stat_c2:
        st.metric("Published", published_count)
    with stat_c3:
        st.metric("Draft", draft_count)
    with stat_c4:
        st.metric("Avg Word Count", avg_words)

    st.divider()

    # ── Search + filters ──────────────────────────────────────────────────────
    filter_col1, filter_col2, filter_col3 = st.columns([3, 1.5, 1.5])
    with filter_col1:
        search_query = st.text_input(
            "Search blogs",
            key="blog_lib_search",
            placeholder="Search by title, keyword, topic...",
            label_visibility="collapsed",
        )
    with filter_col2:
        status_filter = st.selectbox(
            "Status",
            options=["All", "draft", "ready", "published"],
            key="blog_lib_status_filter",
            label_visibility="collapsed",
        )
    with filter_col3:
        type_options = ["All types"] + [
            "Industry Commentary", "Harm Reduction", "Product Education",
            "Campaign Story", "Behind The Scenes", "News-Jack", "Opinion / Thought Leadership",
        ]
        type_filter = st.selectbox(
            "Type",
            options=type_options,
            key="blog_lib_type_filter",
            label_visibility="collapsed",
        )

    # Apply filters
    if search_query.strip():
        display_blogs = search_blogs(search_query.strip())
    else:
        display_blogs = get_all_blogs()

    if status_filter != "All":
        display_blogs = [b for b in display_blogs if b.get("status") == status_filter]

    if type_filter != "All types":
        display_blogs = [b for b in display_blogs if b.get("blog_type") == type_filter]

    if not display_blogs:
        st.info("No blogs found. Write your first post in the Write tab.")
    else:
        st.caption(f"Showing {len(display_blogs)} blog{'s' if len(display_blogs) != 1 else ''}")
        st.divider()

        for blog in display_blogs:
            blog_id = blog.get("id", "")
            title = blog.get("title", "Untitled")
            status = blog.get("status", "draft")
            created_raw = blog.get("created_at", "")
            created_display = created_raw[:10] if created_raw else "Unknown date"
            badge = STATUS_BADGES.get(status, "⚪ Draft")

            with st.expander(
                f"{badge}  |  {title}  —  {created_display}",
                expanded=False,
            ):
                # Meta row
                meta_c1, meta_c2, meta_c3, meta_c4 = st.columns(4)
                with meta_c1:
                    st.caption(f"**Type:** {blog.get('blog_type', '—')}")
                with meta_c2:
                    st.caption(f"**Keyword:** {blog.get('primary_keyword', '—')}")
                with meta_c3:
                    st.caption(f"**Words:** {blog.get('word_count', 0):,}")
                with meta_c4:
                    st.caption(f"**Created:** {created_display}")

                content_tab, actions_tab = st.tabs(["📄 Content", "⚙️ Actions"])

                with content_tab:
                    blog_sections = blog.get("sections", {})
                    if not blog_sections:
                        st.caption("No content stored.")
                    else:
                        for sec_name, sec_content in blog_sections.items():
                            st.markdown(f"**{sec_name}**")
                            st.markdown(sec_content)
                            st.code(sec_content, language=None)
                            st.divider()

                with actions_tab:
                    actions_left, actions_right = st.columns(2)

                    with actions_left:
                        new_status = st.selectbox(
                            "Update status",
                            options=["draft", "ready", "published"],
                            index=["draft", "ready", "published"].index(status),
                            key=f"blog_lib_status_{blog_id}",
                        )
                        if st.button(
                            "Update status",
                            key=f"blog_lib_update_{blog_id}",
                            use_container_width=True,
                        ):
                            try:
                                update_blog_status(blog_id, new_status)
                                st.success(f"Status updated to {new_status}.")
                                st.rerun()
                            except Exception as e:
                                st.error(f"Update failed: {e}")

                        if status != "published":
                            if st.button(
                                "🌐 Mark as Published",
                                key=f"blog_lib_publish_{blog_id}",
                                use_container_width=True,
                            ):
                                try:
                                    update_blog_status(blog_id, "published")
                                    st.success("Marked as published.")
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"Update failed: {e}")

                    with actions_right:
                        st.caption("Danger zone")
                        confirm_key = f"blog_lib_confirm_delete_{blog_id}"
                        if not st.session_state.get(confirm_key):
                            if st.button(
                                "🗑️ Delete",
                                key=f"blog_lib_delete_{blog_id}",
                                use_container_width=True,
                            ):
                                st.session_state[confirm_key] = True
                                st.rerun()
                        else:
                            st.warning("Are you sure? This cannot be undone.")
                            confirm_c1, confirm_c2 = st.columns(2)
                            with confirm_c1:
                                if st.button(
                                    "Yes, delete",
                                    key=f"blog_lib_delete_confirm_{blog_id}",
                                    use_container_width=True,
                                    type="primary",
                                ):
                                    try:
                                        delete_blog(blog_id)
                                        st.success("Blog deleted.")
                                        st.rerun()
                                    except Exception as e:
                                        st.error(f"Delete failed: {e}")
                            with confirm_c2:
                                if st.button(
                                    "Cancel",
                                    key=f"blog_lib_delete_cancel_{blog_id}",
                                    use_container_width=True,
                                ):
                                    st.session_state[confirm_key] = False
                                    st.rerun()
