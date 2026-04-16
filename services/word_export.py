"""
Word document (.docx) export for PR packs.
Creates a formatted Word document that can be opened in Google Docs, Word, or Pages.
No API keys required — pure local generation.
"""

import io
from datetime import datetime


def export_pr_pack_to_docx(pack: dict) -> bytes:
    """
    Export a PR pack to a .docx file.
    Returns the file as bytes (for Streamlit download_button).
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # --- Styles ---
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- Title block ---
    title_para = doc.add_heading("RIOT PR DESK — DRAFT", 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    pack_title = doc.add_heading(pack.get("title", "PR Pack"), 1)

    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {pack.get('created_at', '')[:10]}  |  "
                 f"Position: {pack.get('position_name', '')}  |  "
                 f"Spokesperson: {pack.get('spokesperson_key', '')}  |  "
                 f"Status: {pack.get('status', 'draft').upper()}")
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()  # spacer

    # --- Input context ---
    input_content = pack.get("input_content", "")
    if input_content:
        doc.add_heading("Story Input", 2)
        p = doc.add_paragraph(input_content[:500])
        p.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        doc.add_paragraph()

    # --- Coverage summary ---
    coverage = pack.get("coverage", [])
    if coverage:
        doc.add_heading("Coverage Logged", 2)
        for c in coverage:
            doc.add_paragraph(
                f"• {c.get('publication', '?')} — {c.get('journalist', '')} — "
                f"{c.get('sentiment', '').title()} — reach: {c.get('reach_estimate', 0):,}",
                style="List Bullet"
            )
        doc.add_paragraph()

    # --- PR sections ---
    section_order = [
        "Press Release", "Journalist Pitch Email", "LinkedIn Post",
        "Retailer WhatsApp Comms", "Consumer Social Media Comms", "Internal Briefing",
    ]
    sections = pack.get("sections", {})
    ordered_keys = [k for k in section_order if k in sections]
    ordered_keys += [k for k in sections if k not in section_order]

    for section_name in ordered_keys:
        content = sections[section_name]
        if content.strip().upper().startswith("NOT APPLICABLE"):
            continue

        doc.add_page_break()
        doc.add_heading(section_name, 2)

        # DRAFT watermark line
        draft_para = doc.add_paragraph("⚠️  DRAFT — Requires approval before external use")
        draft_para.runs[0].font.size = Pt(9)
        draft_para.runs[0].font.color.rgb = RGBColor(0xFF, 0x66, 0x00)
        draft_para.runs[0].bold = True

        doc.add_paragraph()

        # Content - split on double newlines for paragraphs
        for block in content.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            # Detect headers (lines starting with # or **text**)
            if block.startswith("##"):
                doc.add_heading(block.lstrip("# ").strip(), 3)
            elif block.startswith("#"):
                doc.add_heading(block.lstrip("# ").strip(), 3)
            else:
                p = doc.add_paragraph(block)

        doc.add_paragraph()

    # --- Tags and comments ---
    tags = pack.get("tags", [])
    comments = pack.get("comments", [])

    if tags or comments:
        doc.add_page_break()
        if tags:
            doc.add_heading("Tags", 3)
            doc.add_paragraph(" | ".join(tags))
        if comments:
            doc.add_heading("Comments & Approvals", 3)
            for c in comments:
                ts = c.get("created_at", "")[:16].replace("T", " ")
                doc.add_paragraph(
                    f"{c.get('author', '?')} ({c.get('type', 'note').replace('_', ' ').title()}) — "
                    f"{c.get('text', '')} — {ts}",
                    style="List Bullet"
                )

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
